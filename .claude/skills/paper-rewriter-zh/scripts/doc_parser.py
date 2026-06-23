"""
文档解析：Word (.docx) / PDF → 纯文本
依赖：python-docx (Word), PyMuPDF (PDF)
用法：$PY doc_parser.py <文件路径>
"""
import sys
import re
from pathlib import Path


def parse_docx(file_path: str) -> str:
    """解析 Word 文档，提取纯文本"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 缺少 python-docx，请运行: uv pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也解析表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)

    return "\n".join(paragraphs)


def parse_pdf(file_path: str) -> str:
    """解析 PDF 文档，提取纯文本"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("错误: 缺少 PyMuPDF，请运行: uv pip install PyMuPDF", file=sys.stderr)
        sys.exit(1)

    doc = fitz.open(file_path)
    paragraphs = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            paragraphs.append(text.strip())
    doc.close()

    return "\n".join(paragraphs)


def parse_doc(file_path: str) -> str:
    """解析旧版 .doc 格式（需要 antiword 或 textract）"""
    # 尝试用 textract
    try:
        import textract
        text = textract.process(file_path).decode("utf-8")
        return text.strip()
    except ImportError:
        pass

    # 提示用户转换格式
    print("错误: .doc 格式需要额外依赖。建议：", file=sys.stderr)
    print("  1. 用 Word 另存为 .docx", file=sys.stderr)
    print("  2. 或安装 textract: uv pip install textract", file=sys.stderr)
    sys.exit(1)


def parse_document(file_path: str) -> str:
    """根据文件扩展名选择解析器"""
    path = Path(file_path)
    if not path.exists():
        print(f"错误: 文件不存在: {file_path}", file=sys.stderr)
        sys.exit(1)

    ext = path.suffix.lower()
    if ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".doc":
        return parse_doc(file_path)
    elif ext == ".txt":
        return path.read_text(encoding="utf-8").strip()
    else:
        print(f"错误: 不支持的格式: {ext}（支持 .docx, .pdf, .doc, .txt）", file=sys.stderr)
        sys.exit(1)


def clean_extracted_text(text: str) -> str:
    """清理提取的文本：去除多余空行、页眉页脚等"""
    # 去除页码模式（如 "- 1 -"、"第1页"）
    text = re.sub(r'\n\s*[-—]\s*\d+\s*[-—]\s*\n', '\n', text)
    text = re.sub(r'\n\s*第\s*\d+\s*页\s*\n', '\n', text)

    # 去除页眉页脚常见模式
    text = re.sub(r'\n\s*中国知网.*?\n', '\n', text)
    text = re.sub(r'\n\s*CNKI.*?\n', '\n', text)

    # 合并多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python doc_parser.py <文件路径>")
        print("支持: .docx, .pdf, .doc, .txt")
        sys.exit(1)

    file_path = sys.argv[1]
    text = parse_document(file_path)
    text = clean_extracted_text(text)

    print(f"提取字数: {len(text)}")
    print("---")
    print(text)
