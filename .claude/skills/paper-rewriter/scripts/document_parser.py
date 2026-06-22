"""
文档解析脚本
从 docx/pdf 文件中提取文本
"""

from pathlib import Path
import re


def parse_docx(file_path: str) -> dict:
    """
    解析 docx 文件

    参数:
        file_path: docx 文件路径

    返回:
        解析结果
    """
    try:
        import docx
    except ImportError:
        return {"error": "需要安装 python-docx: pip install python-docx"}

    path = Path(file_path)
    if not path.exists():
        return {"error": f"文件不存在: {file_path}"}

    try:
        doc = docx.Document(path)

        # 提取段落
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })

        # 识别章节
        sections = identify_sections(paragraphs)

        # 提取元数据
        metadata = extract_metadata(doc)

        return {
            "success": True,
            "file": str(path),
            "paragraph_count": len(paragraphs),
            "paragraphs": paragraphs,
            "sections": sections,
            "metadata": metadata
        }

    except Exception as e:
        return {"error": f"解析失败: {str(e)}"}


def identify_sections(paragraphs: list) -> dict:
    """
    识别论文章节

    参数:
        paragraphs: 段落列表

    返回:
        章节字典
    """
    sections = {
        "title": None,
        "abstract": None,
        "keywords": None,
        "introduction": [],
        "methods": [],
        "results": [],
        "discussion": [],
        "conclusion": [],
        "references": []
    }

    current_section = None

    for para in paragraphs:
        text = para["text"].lower()

        # 识别标题
        if para["index"] == 0 and not sections["title"]:
            sections["title"] = para["text"]
            continue

        # 识别摘要
        if ("abstract" in text or "摘要" in text) and len(para["text"]) < 200:
            current_section = "abstract"
            # 如果段落本身包含摘要内容
            if len(para["text"]) > 100:
                sections["abstract"] = para["text"]
            continue

        # 识别关键词
        if "keywords" in text and len(para["text"]) < 200:
            current_section = "keywords"
            sections["keywords"] = para["text"]
            continue

        # 识别引言
        if re.match(r'^1\.?\s*introduction', text) or text.startswith("1.introduction"):
            current_section = "introduction"
            continue

        # 识别方法
        if re.match(r'^2\.?\s*(study area|data|methodology)', text):
            current_section = "methods"
            continue

        # 识别结果
        if re.match(r'^3\.?\s*results?', text):
            current_section = "results"
            continue

        # 识别讨论
        if re.match(r'^4\.?\s*discussion', text):
            current_section = "discussion"
            continue

        # 识别结论
        if re.match(r'^5\.?\s*conclusion', text):
            current_section = "conclusion"
            continue

        # 识别参考文献
        if re.match(r'^references?$', text) or text.startswith("references"):
            current_section = "references"
            continue

        # 添加到当前章节
        if current_section and current_section in sections:
            if isinstance(sections[current_section], list):
                sections[current_section].append(para["text"])
            elif sections[current_section] is None:
                sections[current_section] = para["text"]

    # 处理摘要（可能是多段）
    if sections["abstract"] is None:
        # 尝试从开头提取
        for para in paragraphs[:20]:
            text = para["text"].lower()
            # 检查是否是摘要段落
            if "abstract" in text and len(para["text"]) > 100:
                # 提取摘要内容（去掉"Abstract:"前缀）
                abstract_text = para["text"]
                if "abstract" in abstract_text.lower():
                    # 找到"Abstract:"后的内容
                    idx = abstract_text.lower().find("abstract")
                    if idx >= 0:
                        abstract_text = abstract_text[idx + len("abstract"):].strip()
                        # 去掉冒号
                        if abstract_text.startswith(":") or abstract_text.startswith("："):
                            abstract_text = abstract_text[1:].strip()
                sections["abstract"] = abstract_text
                break

    return sections


def extract_metadata(doc) -> dict:
    """
    提取文档元数据

    参数:
        doc: docx 文档对象

    返回:
        元数据字典
    """
    metadata = {}

    # 核心属性
    core = doc.core_properties
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["author"] = core.author
    if core.created:
        metadata["created"] = str(core.created)
    if core.modified:
        metadata["modified"] = str(core.modified)

    # 统计信息
    metadata["paragraph_count"] = len(doc.paragraphs)
    metadata["table_count"] = len(doc.tables)

    return metadata


def get_text_for_rewrite(parse_result: dict, section: str = None) -> str:
    """
    获取用于改写的文本

    参数:
        parse_result: 解析结果
        section: 指定章节 (可选)

    返回:
        文本字符串
    """
    if "error" in parse_result:
        return ""

    if section:
        # 获取指定章节
        section_lower = section.lower()
        if section_lower in parse_result["sections"]:
            content = parse_result["sections"][section_lower]
            if isinstance(content, list):
                return "\n\n".join(content)
            elif content:
                return content
        return ""

    # 获取所有文本
    texts = [p["text"] for p in parse_result["paragraphs"]]
    return "\n\n".join(texts)


def parse_pdf(file_path: str) -> dict:
    """
    解析 PDF 文件 (简单版本)

    参数:
        file_path: PDF 文件路径

    返回:
        解析结果
    """
    try:
        import PyPDF2
    except ImportError:
        return {"error": "需要安装 PyPDF2: pip install PyPDF2"}

    path = Path(file_path)
    if not path.exists():
        return {"error": f"文件不存在: {file_path}"}

    try:
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            paragraphs = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    # 按段落分割
                    for j, para_text in enumerate(text.split('\n\n')):
                        para_text = para_text.strip()
                        if para_text and len(para_text) > 10:
                            paragraphs.append({
                                "index": len(paragraphs),
                                "text": para_text,
                                "page": i + 1,
                                "style": "Normal"
                            })

            # 识别章节
            sections = identify_sections(paragraphs)

            return {
                "success": True,
                "file": str(path),
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraphs,
                "sections": sections,
                "metadata": {
                    "page_count": len(reader.pages)
                }
            }

    except Exception as e:
        return {"error": f"解析失败: {str(e)}"}


def parse_document(file_path: str) -> dict:
    """
    解析文档 (自动识别格式)

    参数:
        file_path: 文件路径

    返回:
        解析结果
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == '.docx':
        return parse_docx(file_path)
    elif suffix == '.pdf':
        return parse_pdf(file_path)
    else:
        return {"error": f"不支持的文件格式: {suffix}"}
